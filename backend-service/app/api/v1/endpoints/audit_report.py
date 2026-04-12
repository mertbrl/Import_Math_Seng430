"""Clinical-grade Model Audit & Safety Certificate — python-docx generator.

POST /certificate/download-docx

Accepts a rich JSON payload from the frontend containing the full pipeline
state (demographics, evaluation metrics, SHAP features, subgroup fairness,
checklist state, bias flags).  No backend cache lookups needed — all data
travels in the request body so the document is 100% reproducible.

Document structure
──────────────────
  Page 1 – Cover / Header
    • Institutional letterhead (thick top border via XML)
    • Full title: "AI Clinical Decision Support — Model Audit & Safety Certificate"
    • Generated date, participant, organization
    • STATUS BADGE: red "DEPLOYMENT BLOCKED" or green "CLEARED FOR PILOT"

  Section 1 – Champion Model Summary (KV table)
  Section 2 – Data Representation (demographics bullet list + mismatch warning)
  Section 3 – Top Feature Influences (Global SHAP / numbered list)
  Section 4 – Subgroup Fairness Analysis (styled table, red text for Sens < 50%)
              + Bias warnings paragraph
  Section 5 – EU AI Act Compliance Checklist (table with ✅ / ❌)
  Sign-off   – 4 signature lines
  Footer     – document ID, confidentiality notice
"""
from __future__ import annotations

import io
import uuid
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

router = APIRouter(prefix="/certificate", tags=["certificate"])

# ─── Color palette ─────────────────────────────────────────────────────────────
_NAVY        = "1E3A5F"
_NAVY_LIGHT  = "EAF0F8"
_RED_BG      = "FEE2E2"
_RED_TEXT    = RGBColor(0xDC, 0x26, 0x26)
_GREEN_TEXT  = RGBColor(0x05, 0x96, 0x69)
_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
_AMBER_BG    = "FEF3C7"
_AMBER_TEXT  = RGBColor(0x92, 0x40, 0x08)
_GREY_BG     = "F8FAFC"


# ─── Request schema ─────────────────────────────────────────────────────────────

class DemographicItem(BaseModel):
    label: str
    training_pct: float | None = None
    population_pct: float | None = None

class SubgroupMetric(BaseModel):
    group: str
    n: int = 0
    accuracy: float | None = None
    sensitivity: float | None = None
    specificity: float | None = None

class ChecklistEntry(BaseModel):
    label: str
    done: bool = False

class AuditReportRequest(BaseModel):
    # Identity
    session_id: str = Field(default="demo-session")
    run_id: str = Field(default="run-1")
    participant: str = Field(default="ML Practitioner")
    organization: str = Field(default="Clinical Institution")

    # Step 2 – dataset demographics (optional)
    demographics: list[DemographicItem] = Field(default_factory=list)

    # Step 5 – champion model + evaluation
    champion_name: str = Field(default="Champion Model")
    model_id: str = Field(default="")
    cv_score: float | None = None
    train_test_gap: float | None = None
    overfitting_risk: str = Field(default="unknown")
    accuracy: float | None = None
    sensitivity: float | None = None
    specificity: float | None = None
    auc: float | None = None

    # Step 6 – top SHAP features
    top_features: list[dict[str, Any]] = Field(default_factory=list)

    # Step 7 – fairness
    bias_detected: bool = False
    bias_threshold: float = 0.10
    subgroup_metrics: list[SubgroupMetric] = Field(default_factory=list)
    fairness_warnings: list[str] = Field(default_factory=list)

    # EU AI Act checklist
    checklist: list[ChecklistEntry] = Field(default_factory=list)


# ─── Low-level Word helpers ────────────────────────────────────────────────────

def _set_cell_shading(cell: Any, hex_color: str) -> None:
    tc = cell._tc  # noqa: SLF001
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell: Any, **kwargs: Any) -> None:
    """Set individual borders on a table cell."""
    tc = cell._tc  # noqa: SLF001
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if side in kwargs:
            el = OxmlElement(f"w:{side}")
            for attr, val in kwargs[side].items():
                el.set(qn(f"w:{attr}"), val)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_run(para: Any, text: str, **kwargs: Any) -> Any:
    """Add a run with optional bold/italic/size/color."""
    run = para.add_run(text)
    if kwargs.get("bold"):
        run.bold = True
    if kwargs.get("italic"):
        run.italic = True
    if kwargs.get("size"):
        run.font.size = Pt(kwargs["size"])
    if kwargs.get("color"):
        run.font.color.rgb = kwargs["color"]
    return run


def _add_thick_top_rule(doc: Document) -> None:
    """Insert a thick navy top border (letterhead effect) using a shaded paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()  # noqa: SLF001
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "thick")
    bottom.set(qn("w:sz"), "36")   # 4.5pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _NAVY)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_heading(doc: Document, number: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{number}  {title.upper()}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    pPr = p._p.get_or_add_pPr()  # noqa: SLF001
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _NAVY)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _pct(value: float | None) -> str:
    if value is None:
        return "—"
    return f"{value * 100:.1f}%"


def _spacer(doc: Document, lines: int = 1) -> None:
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


# ─── Section builders ──────────────────────────────────────────────────────────

def _build_cover(doc: Document, payload: AuditReportRequest, doc_id: str) -> None:
    _add_thick_top_rule(doc)

    # Kicker
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_run(kicker, "CONFIDENTIAL — FOR INTERNAL CLINICAL AUDIT USE ONLY",
              bold=True, size=8, color=RGBColor(0x64, 0x74, 0x8B))

    _spacer(doc)

    # Main title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_run(title_p, "AI Clinical Decision Support", bold=True, size=20,
              color=RGBColor(0x1E, 0x3A, 0x5F))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_run(sub_p, "Model Audit & Safety Certificate", bold=True, size=14,
              color=RGBColor(0x1E, 0x3A, 0x5F))

    _spacer(doc)

    # Metadata row
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now = datetime.utcnow()
    _para_run(meta_p,
              f"Generated: {now.strftime('%d %B %Y  %H:%M UTC')}   |   "
              f"Document ID: {doc_id}   |   Session: {payload.session_id}",
              italic=True, size=9, color=RGBColor(0x64, 0x74, 0x8B))

    _spacer(doc)

    # ── STATUS BADGE ──────────────────────────────────────────────────────────
    status_p = doc.add_paragraph()
    status_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_p.paragraph_format.space_before = Pt(6)
    status_p.paragraph_format.space_after = Pt(6)

    if payload.bias_detected:
        badge_text = "⛔  STATUS: DEPLOYMENT BLOCKED — CLINICAL REVIEW REQUIRED"
        badge_color = _RED_TEXT
    else:
        badge_text = "✅  STATUS: CLEARED FOR PILOT DEPLOYMENT"
        badge_color = _GREEN_TEXT

    badge_run = status_p.add_run(badge_text)
    badge_run.bold = True
    badge_run.font.size = Pt(13)
    badge_run.font.color.rgb = badge_color

    # Thin rule after badge
    _add_thick_top_rule(doc)

    _spacer(doc)

    # Participant block
    part_table = doc.add_table(rows=2, cols=2)
    part_table.style = "Table Grid"
    labels = ["Participant / Analyst", "Organization / Institution",
              "Champion Model", "Risk Classification"]
    values = [payload.participant, payload.organization,
              payload.champion_name, payload.overfitting_risk.capitalize() + " Risk"]
    for i, (lbl, val) in enumerate(zip(labels, values)):
        r, c = divmod(i, 2)
        cell = part_table.cell(r, c)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        _para_run(p, lbl + "\n", bold=True, size=8, color=RGBColor(0x64, 0x74, 0x8B))
        _para_run(p, val, size=10)
        _set_cell_shading(cell, _GREY_BG)

    _spacer(doc)


def _build_model_summary(doc: Document, payload: AuditReportRequest) -> None:
    _add_section_heading(doc, "1.", "Champion Model Summary")

    rows = [
        ("Model Algorithm", payload.champion_name),
        ("Model ID", payload.model_id or "—"),
        ("Cross-Validation Score", _pct(payload.cv_score)),
        ("Train → Test Gap", _pct(payload.train_test_gap)),
        ("Overfitting Risk", payload.overfitting_risk.capitalize()),
        ("Test Accuracy", _pct(payload.accuracy)),
        ("Test Sensitivity (Recall)", _pct(payload.sensitivity)),
        ("Test Specificity", _pct(payload.specificity)),
        ("AUC-ROC", _pct(payload.auc)),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        lc = table.cell(i, 0)
        vc = table.cell(i, 1)
        lc.text = key
        lc.paragraphs[0].runs[0].bold = True
        _set_cell_shading(lc, _NAVY_LIGHT)
        vc.text = val
        # Highlight poor sensitivity
        if "Sensitivity" in key and payload.sensitivity is not None and payload.sensitivity < 0.50:
            vc.paragraphs[0].runs[0].font.color.rgb = _RED_TEXT
            vc.paragraphs[0].runs[0].bold = True
            _set_cell_shading(vc, _RED_BG)


def _build_data_representation(doc: Document, payload: AuditReportRequest) -> None:
    _add_section_heading(doc, "2.", "Data Representation & Demographics")

    if not payload.demographics:
        doc.add_paragraph(
            "No demographic data provided. Attach Step 2 dataset profile to enable this section.",
            style="List Bullet",
        ).runs[0].italic = True
        return

    mismatch_items: list[str] = []

    for item in payload.demographics:
        p = doc.add_paragraph(style="List Bullet")
        pct_train = f"{item.training_pct * 100:.1f}%" if item.training_pct is not None else "N/A"
        pct_pop   = f"{item.population_pct * 100:.1f}%" if item.population_pct is not None else "N/A"
        _para_run(p, f"{item.label}:  ", bold=True, size=10)
        _para_run(p, f"Training data = {pct_train}   |   Reference population = {pct_pop}", size=10)

        # Detect >10% mismatch
        if (item.training_pct is not None and item.population_pct is not None
                and abs(item.training_pct - item.population_pct) > 0.10):
            mismatch_items.append(
                f"{item.label}: training {pct_train} vs population {pct_pop} "
                f"(gap = {abs(item.training_pct - item.population_pct) * 100:.1f}pp)"
            )

    if mismatch_items:
        _spacer(doc)
        warn_p = doc.add_paragraph()
        _para_run(warn_p, "⚠  Representativeness Warning: ", bold=True, size=10, color=_AMBER_TEXT)
        _para_run(warn_p,
                  "The following demographic groups have a >10% gap between training data "
                  "and the reference population. This may introduce systematic bias.",
                  size=10, color=_AMBER_TEXT)
        for m in mismatch_items:
            mp = doc.add_paragraph(style="List Bullet")
            _para_run(mp, m, bold=True, size=10, color=_AMBER_TEXT)


def _build_top_features(doc: Document, payload: AuditReportRequest) -> None:
    _add_section_heading(doc, "3.", "Top Feature Influences (Global SHAP)")

    if not payload.top_features:
        doc.add_paragraph("Global SHAP feature data not available. "
                          "Run Step 6 (Explainability) to populate this section.").runs[0].italic = True
        return

    features = payload.top_features[:8]
    for i, feat in enumerate(features, 1):
        name = str(feat.get("feature", "?"))
        imp  = float(feat.get("importance", 0.0))
        p = doc.add_paragraph(style="List Number")
        _para_run(p, f"{name}", bold=True, size=10)
        _para_run(p, f"  —  SHAP importance: {imp:.4f}", size=10,
                  color=RGBColor(0x64, 0x74, 0x8B))


def _build_subgroup_table(doc: Document, payload: AuditReportRequest) -> None:
    _add_section_heading(doc, "4.", "Subgroup Fairness Analysis")

    if not payload.subgroup_metrics:
        doc.add_paragraph("No demographic subgroups detected. The dataset may not contain "
                          "recognised demographic columns (sex, gender, age, race, ethnicity)."
                          ).runs[0].italic = True
        if payload.bias_detected:
            warn_p = doc.add_paragraph()
            _para_run(warn_p, "⚠  Bias flag is still raised from the API response. "
                              "Manual review recommended.", bold=True, size=10, color=_RED_TEXT)
        return

    # Compute best sensitivity for flagging
    sens_vals = [m.sensitivity for m in payload.subgroup_metrics if m.sensitivity is not None]
    best_sens = max(sens_vals) if sens_vals else 1.0

    headers = ["Subgroup", "N", "Accuracy", "Sensitivity ★", "Specificity"]
    col_widths = [Cm(6.5), Cm(1.5), Cm(2.2), Cm(2.8), Cm(2.5)]

    table = doc.add_table(rows=1 + len(payload.subgroup_metrics), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for j, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[j].width = w

    # Header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_run(p, h, bold=True, size=9, color=_WHITE)
        _set_cell_shading(cell, _NAVY)

    # Data rows
    for i, m in enumerate(payload.subgroup_metrics):
        row_idx = i + 1
        values = [m.group, str(m.n), _pct(m.accuracy), _pct(m.sensitivity), _pct(m.specificity)]

        for j, val in enumerate(values):
            cell = table.cell(row_idx, j)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER

            sens = m.sensitivity
            # Sensitivity column: red if < 50% absolute OR degraded vs best by threshold
            is_poor_sensitivity = (j == 3 and sens is not None
                                   and (sens < 0.50 or (best_sens - sens) > payload.bias_threshold))

            if is_poor_sensitivity:
                _para_run(p, val, bold=True, size=9, color=_RED_TEXT)
                _set_cell_shading(cell, _RED_BG)
            elif i % 2 == 0:
                _para_run(p, val, size=9)
                _set_cell_shading(cell, _GREY_BG)
            else:
                _para_run(p, val, size=9)

    # Footnote
    fn_p = doc.add_paragraph()
    fn_p.paragraph_format.space_before = Pt(4)
    _para_run(fn_p, "★ Sensitivity (True Positive Rate) is the primary fairness metric in high-risk clinical AI. "
                    "Red = Sensitivity < 50% or gap >"
                    f" {int(payload.bias_threshold * 100)}pp vs. best performing group.",
              italic=True, size=8, color=RGBColor(0x64, 0x74, 0x8B))

    # Bias warnings
    if payload.fairness_warnings:
        _spacer(doc)
        for w in payload.fairness_warnings:
            wp = doc.add_paragraph(style="List Bullet")
            _para_run(wp, w, bold=True, size=9, color=_RED_TEXT)


def _build_checklist(doc: Document, payload: AuditReportRequest) -> None:
    _add_section_heading(doc, "5.", "EU AI Act Compliance Checklist")

    items = payload.checklist
    if not items:
        # Default checklist if frontend sends nothing
        items = [
            ChecklistEntry(label="Training data documented and version-controlled", done=False),
            ChecklistEntry(label="Performance metrics computed on held-out test set", done=False),
            ChecklistEntry(label="Model trained, validated, and champion selected", done=False),
            ChecklistEntry(label="Model is explainable (SHAP global & local)", done=False),
            ChecklistEntry(label="Subgroup fairness audit completed", done=False),
            ChecklistEntry(label="Human oversight plan approved by clinical lead", done=False),
            ChecklistEntry(label="Risk classification assigned (EU AI Act Art. 6)", done=False),
            ChecklistEntry(label="Incident reporting procedure documented", done=False),
            ChecklistEntry(label="Post-deployment monitoring plan defined", done=False),
            ChecklistEntry(label="Patient data anonymised / pseudonymised (GDPR Art. 4(5))", done=False),
        ]

    completed = sum(1 for it in items if it.done)

    table = doc.add_table(rows=1 + len(items), cols=2)
    table.style = "Table Grid"

    # Header
    for j, h in enumerate(["EU AI Act Compliance Item", "Status"]):
        cell = table.cell(0, j)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        _para_run(p, h, bold=True, size=9, color=_WHITE)
        _set_cell_shading(cell, _NAVY)

    # Items
    for i, item in enumerate(items):
        lc = table.cell(i + 1, 0)
        sc = table.cell(i + 1, 1)
        lc.paragraphs[0].clear()
        sc.paragraphs[0].clear()

        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        sc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if i % 2 == 0:
            _set_cell_shading(lc, _GREY_BG)
            _set_cell_shading(sc, _GREY_BG)

        _para_run(lc.paragraphs[0], item.label, size=9)

        if item.done:
            _para_run(sc.paragraphs[0], "✅  Complete", bold=True, size=9, color=_GREEN_TEXT)
        else:
            _para_run(sc.paragraphs[0], "❌  Pending", bold=True, size=9, color=_RED_TEXT)

    # Summary line
    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_before = Pt(6)
    _para_run(summary_p,
              f"Compliance Progress: {completed} of {len(items)} items completed.",
              bold=True, size=10,
              color=_GREEN_TEXT if completed == len(items) else _RED_TEXT)


def _build_signoff(doc: Document, payload: AuditReportRequest) -> None:
    _add_section_heading(doc, "6.", "Authorisation & Sign-Off")

    intro_p = doc.add_paragraph()
    _para_run(intro_p,
              "The undersigned confirm that this model has undergone the audit process described in this "
              "certificate and that the findings have been reviewed prior to any deployment decision.",
              size=10)

    _spacer(doc, 2)

    sign_rows = [
        ("Clinical Lead / Principal Investigator", ""),
        ("Data Science / ML Lead", ""),
        ("Data Protection Officer (DPO)", ""),
        ("Institutional Review Board Representative", ""),
    ]

    for role, _ in sign_rows:
        p = doc.add_paragraph()
        _para_run(p, f"{role}:  ", bold=True, size=10)
        _para_run(p, "_" * 42 + "     Date:  __________ / __________ / __________",
                  size=10, color=RGBColor(0x9C, 0xA3, 0xAF))
        p.paragraph_format.space_after = Pt(14)


def _build_footer_note(doc: Document, doc_id: str) -> None:
    _add_thick_top_rule(doc)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_run(fp,
              f"Generated by ImportMath ML Platform  •  Document ID: {doc_id}  •  "
              f"Classification: CONFIDENTIAL  •  "
              + datetime.utcnow().strftime("%d %B %Y"),
              italic=True, size=8, color=RGBColor(0x64, 0x74, 0x8B))

    fp2 = doc.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_run(fp2,
              "This document is auto-generated and intended solely for internal clinical audit purposes. "
              "Unauthorised distribution is prohibited.",
              italic=True, size=7, color=RGBColor(0x9C, 0xA3, 0xAF))


# ─── Master document builder ───────────────────────────────────────────────────

def _build_document(payload: AuditReportRequest) -> io.BytesIO:
    doc = Document()

    # Narrow margins for more table real-estate
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    doc_id = str(uuid.uuid4())[:8].upper()

    _build_cover(doc, payload, doc_id)
    _build_model_summary(doc, payload)
    _build_data_representation(doc, payload)
    _build_top_features(doc, payload)
    _build_subgroup_table(doc, payload)
    _build_checklist(doc, payload)
    _build_signoff(doc, payload)
    _build_footer_note(doc, doc_id)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── Endpoint ──────────────────────────────────────────────────────────────────

@router.post("/download-docx")
def download_audit_docx(payload: AuditReportRequest) -> StreamingResponse:
    """Generate and stream the clinical audit certificate as a .docx file."""
    docx_buf = _build_document(payload)
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M")
    filename = f"audit_certificate_{payload.session_id}_{ts}.docx"
    return StreamingResponse(
        docx_buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
